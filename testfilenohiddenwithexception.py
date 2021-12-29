import os
import sys
import docx



doc = docx.Document()
def list_files(startpath):
    for root, dirs, files in os.walk(startpath):
        level = root.replace(startpath, '').count(os.sep)
        # Add a heading of level 0 (Also called Title)
        doc.add_heading(root, level)
        indent = ' ' * 4 * (level)
        print('{}{}/'.format(indent, os.path.basename(root)))
        subindent = ' ' * 4 * (level + 1)
        for f in files:
            if not f.startswith('.'):
                try:
                    doc.add_heading(f, level)
                    print('{}{}'.format(subindent, f))
                except Exception as e:
                    print (e)
                    continue
if __name__=="__main__":
    list_files(sys.args[1])
    doc.save('1.docx')
